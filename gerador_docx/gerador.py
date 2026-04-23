"""Gerador de .docx forense.

Portado 1-para-1 do `sentenca.js` original: mesma fonte, mesmos recuos,
mesma margem de página A4, mesma estrutura de seções e fecho com
assinatura.

Lê um arquivo Python (o `sentenca.py` que o Claude escreve) num sandbox
controlado, executando as funções bp/cp/sh/ch/cc/el do DSL, e monta o
documento Word.

Uso:
    python gerador.py --input sentenca.py --output sentenca.docx
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))

from dsl import Bloco, bp, cc, ch, cp, el, sh  # noqa: E402


# ---------------------------------------------------------------------------
# Constantes de formatação (espelham sentenca.js)
# ---------------------------------------------------------------------------

FONTE = "Times New Roman"
TAMANHO = Pt(12)
RECUO_PRIMEIRA_LINHA = Cm(2.5)
RECUO_CITACAO = Cm(2.5)
ESPACAMENTO = 1.5

# Margens A4 do original: 3cm esq/sup, 2cm dir/inf
MARGEM_SUPERIOR = Cm(3.0)
MARGEM_INFERIOR = Cm(2.0)
MARGEM_ESQUERDA = Cm(3.0)
MARGEM_DIREITA = Cm(2.0)

# Espaçamento antes/depois de citações e subcabeçalhos (em pontos)
ESPACO_ANTES_CITACAO = Pt(6)
ESPACO_APOS_CITACAO = Pt(6)
ESPACO_ANTES_SUBCABECALHO = Pt(14)


# ---------------------------------------------------------------------------
# Sandbox para leitura do arquivo de sentença
# ---------------------------------------------------------------------------

def carregar_sentenca(caminho: Path) -> dict:
    """Executa o .py num namespace controlado e devolve os símbolos."""
    fonte = caminho.read_text(encoding="utf-8")

    ns: dict = {
        "bp": bp,
        "cp": cp,
        "sh": sh,
        "ch": ch,
        "cc": cc,
        "el": el,
        "__builtins__": {"len": len, "range": range, "str": str},
    }

    compilado = compile(fonte, str(caminho), "exec")
    exec(compilado, ns)

    obrigatorios = ("PROCESSO", "RELATORIO", "FUNDAMENTACAO", "DISPOSITIVO")
    faltantes = [n for n in obrigatorios if n not in ns]
    if faltantes:
        raise ValueError(
            f"{caminho} não define: {', '.join(faltantes)}"
        )

    return {nome: ns[nome] for nome in obrigatorios}


# ---------------------------------------------------------------------------
# Renderização de blocos
# ---------------------------------------------------------------------------

def _aplicar_fonte_padrao(paragrafo) -> None:
    for run in paragrafo.runs:
        if run.font.name is None:
            run.font.name = FONTE
        if run.font.size is None:
            run.font.size = TAMANHO


def adicionar_bloco(doc, bloco: Bloco) -> None:
    tipo = bloco.tipo

    if tipo == "el":
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        return

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = ESPACAMENTO

    if tipo == "bp":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = RECUO_PRIMEIRA_LINHA
        p.add_run(bloco.texto)

    elif tipo == "cp":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = RECUO_CITACAO
        p.paragraph_format.space_before = ESPACO_ANTES_CITACAO
        p.paragraph_format.space_after = ESPACO_APOS_CITACAO
        run = p.add_run(bloco.texto)
        run.italic = True

    elif tipo == "sh":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = ESPACO_ANTES_SUBCABECALHO
        run = p.add_run(bloco.texto)
        run.bold = True

    elif tipo == "ch":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(bloco.texto)
        run.bold = True

    elif tipo == "cc":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(bloco.texto)

    else:
        raise ValueError(f"tipo de bloco desconhecido: {tipo}")

    _aplicar_fonte_padrao(p)


# ---------------------------------------------------------------------------
# Montagem do documento
# ---------------------------------------------------------------------------

def gerar(dados: dict, destino: Path, cidade: str = "São Paulo") -> None:
    doc = Document()

    # Margens A4
    for section in doc.sections:
        section.top_margin = MARGEM_SUPERIOR
        section.bottom_margin = MARGEM_INFERIOR
        section.left_margin = MARGEM_ESQUERDA
        section.right_margin = MARGEM_DIREITA

    # Fonte padrão do estilo Normal (pega parágrafos vazios também)
    style = doc.styles["Normal"]
    style.font.name = FONTE
    style.font.size = TAMANHO

    # Ordem exata do sentenca.js:
    # cc(processo), el, ch(SENTENÇA), el, ch(RELATÓRIO), el, ...relatório,
    # el, ch(FUNDAMENTAÇÃO), el, ...fund, el, ch(DISPOSITIVO), el,
    # ...disp, el, el, cc(cidade + data), el, el, cc(___), cc(Juiz(a))

    sequencia = [
        cc(dados["PROCESSO"]),
        el(),
        ch("SENTENÇA"),
        el(),
        ch("RELATÓRIO"),
        el(),
        *dados["RELATORIO"],
        el(),
        ch("FUNDAMENTAÇÃO"),
        el(),
        *dados["FUNDAMENTACAO"],
        el(),
        ch("DISPOSITIVO"),
        el(),
        *dados["DISPOSITIVO"],
        el(),
        el(),
        cc(f"{cidade}, data da assinatura."),
        el(),
        el(),
        cc("__________________________________________________"),
        cc("Juiz(a) de Direito"),
    ]

    for bloco in sequencia:
        adicionar_bloco(doc, bloco)

    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destino)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> int:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--input", required=True, type=Path)
    p.add_argument("--output", required=True, type=Path)
    p.add_argument("--cidade", default="São Paulo")
    args = p.parse_args()

    if not args.input.exists():
        print(f"erro: {args.input} não existe", file=sys.stderr)
        return 1

    dados = carregar_sentenca(args.input)
    gerar(dados, args.output, cidade=args.cidade)
    print(f"Gerado: {args.output}", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
